import os
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import RGBColor, Pt
import requests

def read_paths_from_word(file_path):
    doc = Document(file_path)
    return [p.text for p in doc.paragraphs if p.text]

def process_tag(tag, doc, paragraph=None):
    if isinstance(tag, NavigableString):
        text = tag.strip()
        if text:
            if paragraph is None:
                paragraph = doc.add_paragraph()
            paragraph.add_run(text)
    elif tag.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(tag.get_text())
        run.bold = True
        run.italic = True
        run.font.size = Pt(14)
    elif tag.name == 'div' and 'class' in tag.attrs and ('cell_input docutils container' in tag.attrs['class'] or 'cell_output docutils container' in tag.attrs['class']):
        paragraph = doc.add_paragraph()
        for child in tag.children:
            if isinstance(child, NavigableString):
                run = paragraph.add_run(child.strip())
                run.font.highlight_color = RGBColor(220, 220, 220)
                run.italic = True
            elif isinstance(child, Tag):
                process_tag(child, doc, paragraph)
    elif tag.name == 'p':
        paragraph = doc.add_paragraph()
        for child in tag.children:
            process_tag(child, doc, paragraph)
    elif isinstance(tag, Tag):
        for child in tag.children:
            process_tag(child, doc)

def get_article_content(url, doc):
    response = requests.get(url, verify=False)
    soup = BeautifulSoup(response.content, 'html.parser')
    article_content = soup.find('article', {'class': 'bd-article', 'role': 'main'})
    if article_content:
        process_tag(article_content, doc)

def main():
    host = input("请输入主机名（例如，https://example.com）: ")
    paths_file = input("请输入包含网页路径的Word文件地址: ")
    output_file = input("请输入本地Word文件地址，用于保存抓取的文本: ")

    # 检查输出文件是否已经存在，如果存在则打开，否则创建新的文档
    if os.path.exists(output_file):
        doc = Document(output_file)
    else:
        doc = Document()

    paths = read_paths_from_word(paths_file)

    for path in paths:
        try:
            url = host + path
            get_article_content(url, doc)
        except Exception as e:
            print(f"在处理路径{path}时发生错误: {e}")
            doc.save(output_file)
            break

    doc.save(output_file)

if __name__ == '__main__':
    main()
